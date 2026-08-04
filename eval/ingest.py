"""Full-corpus ingestion for evaluation.

Path A's notebooks loaded only the clean markdown. Real evaluation needs every
format the corpus actually contains, so this applies the parsing ladder across
all of them: PDF and DOCX and XLSX text extraction (P0-P2) and OCR for the
scanned page (P3). Figures (PNG) are skipped — they need P4 captioning (a vision
model), which is not built; questions that depend on them are expected to fail,
honestly, until P4 exists.

Content comes only from corpus/rendered/. Metadata (service, severity, codes)
comes from the source front matter as a manifest — the same split the notebooks
use (see learning/_shared.read_front_matter).
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "learning"))
from _shared import read_front_matter, chunk_fixed, load_tickets   # noqa: E402

RENDERED = ROOT / "corpus" / "rendered"
SOURCE = ROOT / "corpus" / "source"

_OCR = None      # lazy: only built if a scanned PDF needs it


def _pdf_text(path):
    import fitz
    return "".join(p.get_text() for p in fitz.open(path))


def _pdf_ocr(path):
    global _OCR
    import fitz, numpy as np
    if _OCR is None:
        import easyocr
        _OCR = easyocr.Reader(["vi", "en"], gpu=True, verbose=False)
    parts = []
    for page in fitz.open(path):
        pix = page.get_pixmap(dpi=200)
        img = np.frombuffer(pix.samples, dtype="uint8").reshape(pix.height, pix.width, pix.n)
        if pix.n == 4:
            img = img[:, :, :3]
        parts.append("\n".join(_OCR.readtext(img, detail=0, paragraph=True)))
    return "\n".join(parts)


def _docx_text(path):
    """Extract in DOCUMENT ORDER — paragraphs and tables interleaved.

    python-docx's .paragraphs and .tables are separate flat lists; iterating them
    that way appends every table after every paragraph, moving a table away from
    the heading that introduces it. That is a real reading-order defect (it hurts
    retrieval, and it wrecks a naive text-vs-source diff). We walk the body's XML
    children instead so order is preserved.
    """
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    d = docx.Document(path)
    out = []
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            out.append(Paragraph(child, d).text)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, d).rows:
                out.append(" | ".join(c.text for c in row.cells))
    return "\n".join(out)


def _xlsx_text(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _extract(path):
    """Return (text, parse_rung) for one rendered file, or (None, skip-reason)."""
    suffix = path.suffix.lower()
    if suffix == ".md":
        return path.read_text(encoding="utf-8"), "P0"
    if suffix == ".docx":
        return _docx_text(path), "P2"
    if suffix == ".xlsx":
        return _xlsx_text(path), "P2"
    if suffix == ".pdf":
        text = _pdf_text(path)
        if len(text.strip()) < 50:                       # scanned -> OCR (P3)
            return _pdf_ocr(path), "P3"
        return text, "P0"
    if suffix == ".png":
        return None, "needs-P4-captioning"               # figures: no vision model
    return None, "unknown-format"


def load_full_corpus(verbose=True):
    """Every rendered document, parsed by format, chunked, with metadata attached.

    Files sharing a base id (e.g. the multi-sheet D05.en.xlsx + D05.en.sheet02.xlsx)
    are concatenated into one document — otherwise we reproduce the only-first-sheet
    bug (F1.14) in our own ingest. Language versions stay separate doc ids; grouping
    by base doc is the retrieval layer's job (notebook 03).

    Returns (chunks, report) where report maps doc_id -> (parse_rungs, n_chars).
    """
    # group rendered files by base document id (strip a trailing .sheetNN)
    groups = {}
    for path in sorted(RENDERED.glob("*.*")):
        stem = path.name.rsplit(".", 1)[0]               # "D05.en.sheet02" or "D06.vi"
        base_id = stem.split(".sheet")[0]                # -> "D05.en"
        groups.setdefault(base_id, []).append(path)

    chunks, report = [], {}
    for base_id, paths in sorted(groups.items()):
        parts = base_id.split(".")
        doc, lang = parts[0], (parts[1] if len(parts) > 1 else "en")
        meta = read_front_matter(SOURCE / f"{doc}.{lang}.md")
        texts, rungs = [], []
        for p in sorted(paths):
            text, rung = _extract(p)
            if text:
                texts.append(text)
                rungs.append(rung)
        if not texts:
            report[base_id] = (rungs[0] if rungs else "skip", 0)
            continue
        full = "\n".join(texts)
        report[base_id] = ("+".join(sorted(set(rungs))), len(full.strip()))
        for piece in chunk_fixed(full):
            chunks.append({"doc_id": base_id, "doc": doc, "lang": lang,
                           "text": piece, "meta": meta, "kind": "doc"})
    if verbose:
        for k in sorted(report):
            r, n = report[k]
            print(f"  {k:16s} {r:20s} {n:6d} chars")
    return chunks, report


def load_ticket_chunks():
    rows, tchunks = load_tickets()
    for c in tchunks:
        c.update(doc=c["doc_id"], lang="vi", kind="ticket")
    return rows, tchunks


if __name__ == "__main__":
    chunks, report = load_full_corpus()
    print(f"\n{len(chunks)} document chunks from {len(report)} documents")
    scanned = [k for k, (r, _) in report.items() if r == "P3"]
    figures = [k for k, (r, _) in report.items() if r.startswith("needs")]
    print(f"OCR'd (P3): {scanned}")
    print(f"skipped, need P4 captioning: {figures}")
